# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_table(doc, headers, rows, header_color='2E75B6'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], header_color)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10)
    
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for para in row_cells[col_idx].paragraphs:
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    return table

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('OUTDOOR-FIT APP 升级需求可行性分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    add_para(doc, '文档来源：OUTDOOR FIT APP Upgrade.pptx.pdf', size=10)
    add_para(doc, '分析对象：outdoor-rehab-fit-anesthesia 项目代码库', size=10)
    add_para(doc, '生成日期：2026年7月20日', size=10)
    doc.add_paragraph()

    # Section 1: Current Status
    add_heading(doc, '一、当前 App 现状', 1)
    add_para(doc, '这是一个基于 Expo React Native 的离线内容型 App，主要能力如下：')
    
    add_table(doc, ['已有功能', '缺失功能'], [
        ['约 28 种器材介绍（硬编码在 constants/Equipments.js）', '用户登录 / 权限系统'],
        ['音频播放 + YouTube 视频', '后端服务 / API'],
        ['收藏功能（AsyncStorage）', '运动记录、目标设定'],
        ['地图页（WebView 外链）', '群组聊天 / SMS'],
        ['中英文 i18n', '使用数据统计'],
        ['免责声明 + PAR-Q 问卷流程', 'CMS 内容管理'],
    ])
    
    add_para(doc, '当前架构：纯前端 + 本地存储，无服务器。PDF 中约 70% 的功能需要后端支持，这是一次从「内容展示 App」到「研究平台型 App」的架构升级。', bold=True)

    # Section 2: Item by item analysis
    add_heading(doc, '二、逐项可行性分析', 1)

    items = [
        {
            'title': '#1 登录 / 访客模式',
            'feasibility': '可行，中等工作量',
            'requirement': '登录用户全功能；访客只能看部分视频/音频，受限功能弹窗提示。',
            'current': '只有免责声明流程，无账号体系。',
            'path': '引入 Firebase Auth / Supabase Auth；路由守卫 + 功能权限矩阵；参考 PDF 中的「风起航」「傲龄汇」等 App 模式。',
            'risk': '需先和需求方确认「访客可访问哪些内容」的完整清单。',
            'estimate': '2–3 周',
        },
        {
            'title': '#2 器材定位 GIS',
            'feasibility': '可行，依赖数据',
            'requirement': '帮助用户定位特定类型户外健身器材。',
            'current': 'location/index.js 只是 WebView 打开外部地图 URL。',
            'path': '接入 react-native-maps 或 Mapbox；加载 GeoJSON / GIS 点位数据；按器材类型筛选、导航。',
            'risk': 'PDF 写明 GIS 数据后续才提供。',
            'estimate': '3–4 周（数据到位后）',
        },
        {
            'title': '#3 内容管理系统 CMS',
            'feasibility': '可行，架构改动大',
            'requirement': '为现有离线 App 建立内容管理系统。',
            'current': '所有内容硬编码在 Equipments.js（约 1100+ 行），含图片、音频、文案。',
            'path': '选型 Strapi / Sanity / Supabase + Admin Panel；App 改为 API 拉取内容；图片/音频走 CDN；可配合已有 expo-updates 做 OTA。',
            'risk': '这是从离线到在线的核心架构转变，影响面最大。',
            'estimate': '4–6 周',
        },
        {
            'title': '#4 App 使用数据统计',
            'feasibility': '可行，需后端',
            'requirement': '记录 App 使用天数；1.5 小时无操作自动登出；统计器材介绍页打开次数；统计音频播放次数；统计目标达成次数；统计户外/居家训练完成次数。',
            'current': '无任何埋点或统计。',
            'path': '前端 Event Tracking 层；后端聚合（Firebase Analytics + 自定义 API，或 Mixpanel）；1.5h 超时用 AppState + 定时器实现。',
            'risk': '这是研究数据采集，需考虑隐私合规（IRB、数据脱敏、知情同意）。',
            'estimate': '3–4 周',
        },
        {
            'title': '#5 运动 Session 记录',
            'feasibility': '可行，部分功能复杂',
            'requirement': '计时器、暂停/恢复、情绪量表、RPE 量表、运动日志、拍照识别器材、Session 日志、分享。',
            'current': '首页只有 4 个导航按钮，无相关功能。',
            'path': '计时器 UI（低难度）；情绪/RPE 量表（低难度）；Session 摘要页（中难度）；拍照识别器材（高难度，需 CV/ML）；Session 日志列表（中难度）；分享到群组（依赖 #7）。',
            'risk': '「拍照自动识别器材名称」是技术难点。建议分阶段：Phase 1 手动选择 → Phase 2 图像相似度匹配 → Phase 3 完整 CV 识别。',
            'estimate': '5–8 周',
        },
        {
            'title': '#6 目标设定与进度追踪',
            'feasibility': '可行，工作量中等',
            'requirement': '3 类预设目标（有氧/平衡/肌力），可调整每周天数；SMART 目标 + 写作指引；首页周历 + 进度条；每周可延续目标。',
            'current': '首页无目标相关 UI。',
            'path': '目标数据模型 + 本地/云端存储；与 #5 Session 数据联动计算进度；首页大改：加周历、进度条、SMART 目标区块。',
            'risk': '无',
            'estimate': '4–5 周',
        },
        {
            'title': '#7 内置 SMS 群组功能',
            'feasibility': '可行但工作量最大',
            'requirement': '实验组参与者 + 项目人员群组；支持文字、图片、Emoji 反应；Session 分享；鼓励互评。',
            'current': '完全空白。',
            'path': '实时通讯：Firebase Realtime DB / Supabase Realtime / Socket.io；群组管理、消息列表、图片上传；Push Notification；内容 moderation。',
            'risk': '本质是做一个迷你社交 App，复杂度高；需考虑聊天审核、隐私、存储成本；PDF 写「SMS」，实际应是 App 内群组聊天，不是真 SMS。',
            'estimate': '8–12 周',
        },
        {
            'title': '#8 Chatbot',
            'feasibility': '可行，分阶段实现',
            'requirement': '群组 Leader 角色：每周一早上发提醒；问答：基于 FAQ 数据集回答 App 使用问题。',
            'current': '完全空白。',
            'path': '规则 + FAQ 匹配（低复杂度）；定时推送（中复杂度）；LLM GPT 等（中-高复杂度，有 API 成本）。建议先用 FAQ 规则引擎，后续按需升级 LLM。',
            'risk': '无',
            'estimate': '3–5 周',
        },
        {
            'title': '#9 App 内容扩充',
            'feasibility': '可行，依赖内容',
            'requirement': '增加器材数量；每种器材多加 1 个 YouTube 视频。',
            'current': '数据模型已有 youtubeKey 字段，扩展成本低。',
            'path': '扩展数据模型支持多视频字段；内容录入。',
            'risk': '内容后续提供。',
            'estimate': '1–2 周（内容到位 + CMS 就绪后）',
        },
        {
            'title': '#10 学习进度追踪',
            'feasibility': '可行，工作量小',
            'requirement': '追踪用户在「我的收藏」中是否已阅读/观看教育材料。',
            'current': '收藏页只有书签，无阅读状态。',
            'path': '在 detail.js 记录页面打开、视频播放、音频播放；收藏列表显示已读/未读标记；AsyncStorage 或后端存储。',
            'risk': '无',
            'estimate': '1–2 周',
        },
        {
            'title': '#11 文字转语音 TTS',
            'feasibility': '可行，工作量小',
            'requirement': '将文字内容转为语音朗读。',
            'current': '无 TTS 功能。',
            'path': 'expo-speech 或 react-native-tts；在器材详情页加「朗读」按钮；支持中英文。',
            'risk': '无',
            'estimate': '1 周',
        },
    ]

    for item in items:
        add_heading(doc, item['title'], 2)
        add_para(doc, f"可行性：{item['feasibility']}", bold=True)
        add_para(doc, f"需求：{item['requirement']}")
        add_para(doc, f"现状：{item['current']}")
        add_para(doc, f"实现路径：{item['path']}")
        add_para(doc, f"风险/注意：{item['risk']}")
        add_para(doc, f"预估工期：{item['estimate']}")
        doc.add_paragraph()

    # Section 3: Overall assessment
    add_heading(doc, '三、总体评估', 1)
    
    add_heading(doc, '3.1 架构变化', 2)
    add_para(doc, '现有架构：Expo RN App → 硬编码内容 + AsyncStorage（仅本机）')
    add_para(doc, '目标架构：Expo RN App → Auth 层 → Backend API → CMS / Analytics DB / Realtime Chat / Chatbot Service + Local Cache')
    add_para(doc, '这是一次从离线内容 App 到带账号、数据收集、社交互动的研究平台的架构升级。')

    add_heading(doc, '3.2 分期实施建议', 2)
    add_table(doc, ['阶段', '功能', '周期', '依赖'], [
        ['Phase 0', '后端架构选型 + 搭建', '2–3 周', '无'],
        ['Phase 1', '#1 登录、#11 TTS、#10 学习进度', '4–5 周', 'Phase 0'],
        ['Phase 2', '#5 Session 记录、#6 目标设定', '6–8 周', 'Phase 1'],
        ['Phase 3', '#4 数据统计、#3 CMS', '4–6 周', 'Phase 0'],
        ['Phase 4', '#2 GIS、#9 内容扩充', '3–5 周', '客户提供数据'],
        ['Phase 5', '#7 群组聊天、#8 Chatbot', '8–12 周', 'Phase 2'],
    ])
    add_para(doc, '总预估：约 6–9 个月（2–3 名全职开发）', bold=True)

    add_heading(doc, '3.3 主要风险', 2)
    risks = [
        '无后端 — 当前最大缺口，多数功能无法独立在前端完成',
        '拍照识别器材 — 建议分阶段，先做手动选择',
        '群组聊天 — 范围大，需和需求方确认 MVP 范围',
        '研究数据合规 — 使用统计涉及 IRB、隐私政策',
        '内容依赖 — GIS 数据、新器材、新视频、FAQ 数据集都待提供',
        '首页重构 — #5、#6 会大幅改变首页布局',
    ]
    for i, risk in enumerate(risks, 1):
        add_para(doc, f'{i}. {risk}')

    add_heading(doc, '3.4 技术栈建议', 2)
    add_table(doc, ['组件', '推荐方案'], [
        ['后端', 'Supabase（Auth + DB + Realtime + Storage 一体）'],
        ['CMS', 'Strapi 或 Supabase Admin'],
        ['地图', 'react-native-maps + GeoJSON'],
        ['聊天', 'Supabase Realtime 或 Firebase'],
        ['分析', '自定义 API + Supabase'],
        ['TTS', 'expo-speech'],
        ['图像识别', 'Phase 1 手动选择 → Phase 2 Google Vision API'],
    ])

    # Section 4: Summary
    add_heading(doc, '四、总结', 1)
    add_para(doc, '整体可行，但这不是小修小补，而是从离线内容 App 升级为带账号、数据收集、社交互动的研究平台。')
    doc.add_paragraph()
    add_para(doc, '建议下一步：', bold=True)
    next_steps = [
        '和需求方确认 Phase 0 后端方案及预算',
        '明确访客 vs 登录用户的功能权限表',
        '确认 #7 群组聊天的 MVP 范围（是否 Phase 1 就要）',
        '向客户索要 GIS 数据、新内容、FAQ 数据集的时间表',
        '确认研究数据收集的伦理与合规要求',
    ]
    for i, step in enumerate(next_steps, 1):
        add_para(doc, f'{i}. {step}')

    # Save
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'OUTDOOR-FIT_APP_Feasibility_Report.docx')
    doc.save(output_path)
    
    downloads_path = os.path.join(os.path.expanduser('~'), 'Downloads', 'OUTDOOR-FIT_APP_Feasibility_Report.docx')
    doc.save(downloads_path)
    print(f'Report saved to: {output_path}')
    print(f'Also copied to: {downloads_path}')

if __name__ == '__main__':
    main()
